from docx import Document
from .parser import Parser
class DocxParser(Parser):
    def extract_heading(self, file_path: str):
        headings = []
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):  # Check if the style is a heading
                headings.append(para.text.strip())
        return headings

    def extract_content(self, file_path: str):
        content = []
        doc = Document(file_path)
        for para in doc.paragraphs:
            if not para.style.name.startswith('Heading'):  # Skip headings
                content.append(para.text.strip())
        return content

    def parse(self, file_path: str):
        headings = self.extract_heading(file_path)
        content = self.extract_content(file_path)
        return headings, content
